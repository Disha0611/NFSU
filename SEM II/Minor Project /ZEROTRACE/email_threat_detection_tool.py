import os
import json
import base64
import re
import logging
import PyPDF2
from PyPDF2 import PdfReader
import hashlib
from docx import Document
from email.utils import getaddresses
from google.oauth2.credentials import Credentials
from googleapiclient.discovery import build
from email import message_from_bytes
import time

# ========== SETUP ==========
SCOPES = ['https://www.googleapis.com/auth/gmail.readonly']
TOKEN_DIR = 'tokens'
CONFIDENTIAL_KEYWORDS = [
    "zerotrace internal", "zerotrace confidential", "zr-project",
    "zr-financials", "zr-strategy", "employee evaluation", "client contracts",
    "budget breakdown", "nda", "project roadmap", "proprietary algorithm",
    "source code", "performance review", "internal audit", "financial forecast"
]
PHISHING_KEYWORDS = [
    "urgent", "click here", "view as soon as possible", "your account has been compromised",
    "reset your password", "suspended account", "unauthorized access", "important notice",
    "verify your identity", "payment required", "security alert", "login immediately", "www", ".com"
]
ORG_DOMAIN = 'zerotrace@gmail.com'
LOGFILE = 'threat_detection.log'
HISTORY_FILE = 'user_history.json'
# Simple in-memory tracker (replace with file/db for persistent storage)
zip_file_history = {}
PROCESSED_IDS_FILE = "processed.json"
PROCESSED_IDS = set()


# ===== EXPLICIT LOGGER SETUP =====
logger = logging.getLogger("threat_logger")
logger.setLevel(logging.WARNING)

if not logger.handlers:
    file_handler = logging.FileHandler(LOGFILE)
    file_handler.setFormatter(logging.Formatter('%(asctime)s - %(levelname)s - %(message)s'))

    stream_handler = logging.StreamHandler()
    stream_handler.setFormatter(logging.Formatter('%(asctime)s - %(levelname)s - %(message)s'))

    logger.addHandler(file_handler)
    logger.addHandler(stream_handler)

COLORS = {
    "low": "\033[94m",
    "medium": "\033[93m",
    "high": "\033[91m",
    "critical": "\033[41m",
    "reset": "\033[0m"
}

DISPLAYED_ALERTS = set()

with open('user_roles.json') as f:
    USERS = json.load(f)

if os.path.exists(HISTORY_FILE):
    with open(HISTORY_FILE, 'r') as f:
        HISTORY = json.load(f)
else:
    HISTORY = {}

def is_admin(user):
    return USERS.get(user, {}).get('role') == 'admin'

def is_any_external(recipients):
    return any(ORG_DOMAIN not in r for r in recipients if r)


def contains_phishing_keywords(text):
    return any(word.lower() in text.lower() for word in PHISHING_KEYWORDS)

def is_confidential(text):
    return any(word.lower() in text.lower() for word in CONFIDENTIAL_KEYWORDS)

def compute_file_hash(file_path):
    sha256 = hashlib.sha256()
    with open(file_path, 'rb') as f:
        for chunk in iter(lambda: f.read(4096), b""):
            sha256.update(chunk)
    return sha256.hexdigest()

def extract_pdf_content(file_path):
    try:
        with open(file_path, "rb") as file:
            reader = PyPDF2.PdfReader(file)
            text = ""
            for page in reader.pages:
                text += page.extract_text()
            return text
    except Exception as e:
        logger.error(f"Error reading PDF content: {e}")
        return None

def extract_docx_content(file_path):
    try:
        doc = Document(file_path)
        text = ""
        for para in doc.paragraphs:
            text += para.text + "\n"
        return text
    except Exception as e:
        logger.error(f"Error reading DOCX content: {e}")
        return None

def extract_txt_content(file_path):
    try:
        with open(file_path, "r", encoding="utf-8") as file:
            return file.read()
    except Exception as e:
        logger.error(f"Error reading TXT content: {e}")
        return None

def check_txt_for_signature(file_path, signature):
    zero_width_map = {'\u200b': '0', '\u200c': '1'}
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            content = file.read()
            hidden = ''.join(c for c in content if c in zero_width_map)
            if not hidden:
                return False
            binary = ''.join(zero_width_map[c] for c in hidden)
            chars = [chr(int(binary[i:i+8], 2)) for i in range(0, len(binary), 8)]
            extracted_signature = ''.join(chars)
            if signature in extracted_signature:
                print(f"Confidential zero-width signature found in {file_path}")
                return True
    except Exception as e:
        logger.error(f"Error checking signature: {e}")
    return False

def check_docx_for_signature(file_path, expected_signature="zerotrace report - high"):
    doc = Document(file_path)
    return doc.core_properties.subject == expected_signature

def check_pdf_for_signature(file_path, expected_signature="zerotrace report - high"):
    reader = PdfReader(file_path)
    metadata = reader.metadata
    return metadata.get("/ConfidentialSignature") == expected_signature

def extract_pdf_metadata(file_path):
    try:
        reader = PdfReader(file_path)
        metadata = reader.metadata
        return metadata.get('/ConfidentialSignature', None)
    except Exception as e:
        logger.error(f"Error extracting PDF metadata: {e}")
        return None

def download_attachment(service, message_id, attachment_id, file_path):
    try:
        attachment = service.users().messages().attachments().get(
            userId="me", messageId=message_id, id=attachment_id).execute()
        file_data = base64.urlsafe_b64decode(attachment['data'])
        with open(file_path, 'wb') as f:
            f.write(file_data)
        return True
    except Exception as e:
        logger.error(f"Error downloading attachment: {e}")
        return False

def log_alert(username, message, severity="low", intent="accidental"):
    color = COLORS.get(severity, '')
    reset = COLORS["reset"]
    log_line = f"[{severity.upper()}][{intent.upper()}] {username}: {message}"

    # Prevent repeat console display
    if log_line in DISPLAYED_ALERTS:
        return
    DISPLAYED_ALERTS.add(log_line)

    # Print once in console
    print(f"{color}{log_line}{reset}")

    # Log once in file if not already present and if severity warrants
    if severity in ["medium", "high", "critical"]:
        if os.path.exists(LOGFILE):
            with open(LOGFILE, 'r') as log_file:
                if any(log_line in line for line in log_file.readlines()):
                    return  # Already logged
        logger.warning(log_line)

def get_gmail_service(username):
    token_path = os.path.join(TOKEN_DIR, f'token_{username}.json')
    creds = Credentials.from_authorized_user_file(token_path, SCOPES)
    return build('gmail', 'v1', credentials=creds)

def save_history():
    with open(HISTORY_FILE, 'w') as f:
        json.dump(HISTORY, f, indent=2)

def save_processed_ids():
    with open(PROCESSED_IDS_FILE, "w") as f:
        json.dump(list(PROCESSED_IDS), f, indent=2)
        
def update_history(username, flag):
    import inspect

    # Try to auto-detect the message_id from the calling context (if available)
    frame = inspect.currentframe().f_back
    message_id = frame.f_locals.get("message_id")

    # If we have a message_id and it's already processed, skip updating
    if message_id and message_id in PROCESSED_IDS:
        return

    if username not in HISTORY:
        HISTORY[username] = {}
    
    # Split the flag string by comma and count each flag separately
    flags = flag.split(",")
    for flag in flags:
        flag = flag.strip()
        if flag not in HISTORY[username]:
            HISTORY[username][flag] = 0
        HISTORY[username][flag] += 1
        
    # Mark message ID as processed
    if message_id:
        PROCESSED_IDS.add(message_id)
        save_processed_ids()

    save_history()
    
def classify_intention(flags, is_external, username, subject, sender, all_recipients):
    # Helper flags
    is_confidential = "sender_confidential" in flags or "recv_confidential" in flags or "sender_confidential_attachment" in flags or "recv_confidential_attachment" in flags
    is_phishing = "sender_phishing" in flags or "recv_phishing" in flags or "phishing" in flags
    is_normal_attachment = "normal_attachment" in flags
    is_admin = "admin_action" in flags
    is_internal_to_internal = ORG_DOMAIN in sender and all(ORG_DOMAIN in r for r in all_recipients if r.strip())
    is_external_to_internal = (ORG_DOMAIN not in sender and all(ORG_DOMAIN in r for r in all_recipients if r.strip()))
    is_internal_to_external = (ORG_DOMAIN in sender and any(ORG_DOMAIN not in r for r in all_recipients if r.strip()))
    is_sender_internal = ORG_DOMAIN in sender
    

    # Historical behavior
    history = HISTORY.get(username, {})
    repeated_sender_confidential = history.get("sender_confidential", 0) > 1
    repeated_recv_confidential = history.get("recv_confidential", 0) > 1
    repeated_sender_confidential_attachment = history.get("sender_confidential_attachment", 0) > 1
    repeated_recv_confidential_attachment = history.get("recv_confidential_attachment", 0) > 1
    repeated_phishing = history.get("sender_phishing", 0) > 1
    repeated_recv_phishing = history.get("recv_phishing", 0) > 1
    repeated_normal = history.get("normal_attachment", 0) > 1
    repeated_external = history.get("external",0) > 1

    # === RULES ===
    # 1. PHISHING CASES
    if is_phishing:
        if "sender_phishing" in flags:
            if repeated_phishing:   
                return "intentional"
            return "accidental"
        elif repeated_recv_phishing : 
            return "intentional"
        else :
            return "intentional"
    
    if "recv_phishing" in flags and "sender_phishing" in flags and is_sender_internal:
        return "accidental"  # careless forwarding

    if "phishing_obfuscated" in flags:
        if repeated_phishing or is_sender_internal:
            return "intentional"
        return "accidental"
    
    if is_phishing and is_external_to_internal and len(all_recipients) > 3:
        return "intentional"
    
    if is_sender_internal and is_internal_to_internal and len(all_recipients) > 5 and "sender_phishing" in flags:
        return "intentional"

    #if "spoofing_detected" in flags:
     #   return "intentional"
   

    # 4. Normal internal-to-internal emails
    if is_internal_to_internal:
        return "normal"

    # === 2. CONFIDENTIAL ATTACHMENT CASES ===

    # Case A: Sent confidential attachment
    if "sender_confidential_attachment" in flags:
        if is_internal_to_external:
            # First-time case = accidental, repeated = intentional
            return "intentional" if repeated_sender_confidential_attachment else "accidental"
        elif is_internal_to_internal:
            return "normal"

    # Case B: Received confidential attachment
    if "recv_confidential_attachment" in flags:
        if is_external_to_internal:
            # First-time case = accidental, repeated = intentional
            return "intentional" if repeated_recv_confidential_attachment else "accidental"

    # Case D: Obfuscated confidential file received
    if "recv_confidential_attachment" in flags and "phishing_obfuscated" in flags:
        return "intentional"

    # Case E: Internal send with both normal + confidential attachments (mixed behavior)
    if "sender_confidential_attachment" in flags and "normal_attachment" in flags:
        if is_internal_to_external:
            return "intentional" if repeated_sender_confidential_attachment else "accidental"

     
    # === 3. CONFIDENTIAL NON-ATTACHMENT CASES ===
    if "sender_confidential" in flags:
        if is_internal_to_external:
            return "accidental" if repeated_sender_confidential else "accidental"
        elif is_internal_to_internal:
            return "normal"

    if "recv_confidential" in flags:
        if is_external_to_internal:
            return "accidental" if not repeated_recv_confidential else "accidental"



   # 3. Normal files sent externally
    if is_normal_attachment and is_internal_to_external:
        return "intentional" if repeated_external else "accidental"

  
    # 5. Received from external sources
    if is_external_to_internal:
        if "recv_confidential_attachment" in flags or "recv_confidential" in flags:
            return "intentional" if (repeated_recv_confidential or repeated_recv_confidential_attachment) else "accidental"
        elif "external" in flags:
            return "intentional" if repeated_external else "accidental"

    # 6. Sent externally
    if is_internal_to_external:
        if is_confidential or is_normal_attachment:
            return "intentional" if repeated_external else "accidental"
        return "accidental"

    # Default rule
    return "normal"


def attachment_contains_confidential_keywords(service, message_id, parts, signature="zerotrace report - high"):
    for part in parts:
        if 'filename' in part and part['filename']:
            attachment_id = part['body'].get('attachmentId')
            if attachment_id:
                file_path = f"temp_{part['filename']}"
                if download_attachment(service, message_id, attachment_id, file_path):
                    file_extension = os.path.splitext(file_path)[1].lower()
                    file_content = None
                    signature_found = False

                    if file_extension == ".pdf":
                        file_content = extract_pdf_content(file_path)
                        signature_found = check_pdf_for_signature(file_path, signature)
                    elif file_extension == ".docx":
                        file_content = extract_docx_content(file_path)
                        signature_found = check_docx_for_signature(file_path, signature)
                    elif file_extension == ".txt":
                        file_content = extract_txt_content(file_path)
                        signature_found = check_txt_for_signature(file_path, signature)

                    file_hash = compute_file_hash(file_path)
                    os.remove(file_path)

                    if (file_content and is_confidential(file_content)) or signature_found:
                        print(f"[CONFIDENTIAL FILE] Name: {part['filename']}, SHA256: {file_hash}")
                        return True, part['filename'], file_hash, file_content or "CONFIDENTIAL SIGNATURE"
    return False, None, None, None
    

def process_zip_attachment(service, message_id, part, sender, all_recipients, filename_tracker):
    filename = part.get("filename", "")
    attachment_id = part["body"].get("attachmentId")

    if filename.lower().endswith(".zip") and attachment_id:
        temp_path = f"temp_{filename}"
        if download_attachment(service, message_id, attachment_id, temp_path):
            file_hash = compute_file_hash(temp_path)
            os.remove(temp_path)

            # Count hash occurrence
            if file_hash not in filename_tracker:
                filename_tracker[file_hash] = {"count": 1, "all_recipients": [all_recipients]}
            else:
                filename_tracker[file_hash]["count"] += 1
                filename_tracker[file_hash]["all_recipients"].append(all_recipients)

            # Determine type
            is_external = is_any_external(all_recipients)

            if filename_tracker[file_hash]["count"] > 1:
                severity = "high"
                intention = "intentional"
            else:
                severity = "high"
                intention = "accidental"

            print(f"\033[95m[ZIP ALERT][{severity.upper()}][{intention.upper()}] {sender} sent .zip attachment '{filename}' to {all_recipients} (hash: {file_hash})\033[0m")
            return True, severity, intention, filename, file_hash

    return False, None, None, None, None


def extract_emails(header_value):
    """
    Safely extract email addresses from a header (To, Cc, Bcc).
    """

    return [email.lower() for name, email in getaddresses([header_value])]

def get_email_body(payload):
    """Extracts the plain text or HTML body from the email payload."""
    if 'parts' in payload:
        for part in payload['parts']:
            if part['mimeType'] == 'text/plain':
                body_data = part['body'].get('data')
                if body_data:
                    return base64.urlsafe_b64decode(body_data).decode('utf-8', errors='ignore')
            elif part['mimeType'] == 'text/html':
                body_data = part['body'].get('data')
                if body_data:
                    return base64.urlsafe_b64decode(body_data).decode('utf-8', errors='ignore')
    elif payload.get('body', {}).get('data'):
        # Sometimes the message has no parts, just a body
        body_data = payload['body']['data']
        return base64.urlsafe_b64decode(body_data).decode('utf-8', errors='ignore')
    return ""

def detect_obfuscated_phishing(body):
    obfuscated_keywords = [
        r"p[\W_]*a[\W_]*s[\W_]*s[\W_]*w[\W_]*o[\W_]*r[\W_]*d",  # p@ssword, pa_ssword
        r"l[\W_]*o[\W_]*g[\W_]*i[\W_]*n",
        r"v[\W_]*e[\W_]*r[\W_]*i[\W_]*f[\W_]*y",
        r"a[\W_]*c[\W_]*c[\W_]*o[\W_]*u[\W_]*n[\W_]*t"
    ]
    for pattern in obfuscated_keywords:
        if re.search(pattern, body, re.IGNORECASE):
            return True
    return False



def detect_spoofing(sender, headers):
    """
    Detects if the display name or domain is spoofed.
    Examples:
    - Display name = "CEO" but email is from a free domain
    - Email says it's from @zerotrace but the domain is off by 1 char
    """
    from_header = next((h['value'] for h in headers if h['name'] == 'From'), "")
    reply_to = next((h['value'] for h in headers if h['name'] == 'Reply-To'), "")
    
    sender_emails = extract_emails(from_header)
    reply_to_emails = extract_emails(reply_to)

    if not sender_emails:
        return False

    sender_email = sender_emails[0]
    display_name_match = re.match(r"^(.*?)\s*<", from_header)
    display_name = display_name_match.group(1).strip().lower() if display_name_match else ""

    suspicious_names = ["ceo", "admin", "it support", "security", "hr"]
    if display_name in suspicious_names and "zerotrace" not in sender_email:
        return True

    # Example of domain typosquatting detection
    if "zerotrace" in display_name and "zerotrace" not in sender_email:
        return True

    # Mismatch in Reply-To
    if reply_to_emails and reply_to_emails[0] != sender_email:
        return True

    return False


def analyze_email_content(username, message, service, message_id):
    
    payload = message.get("payload", {})
    logger.info(f"Full email payload: {payload}")
    
    label_ids = message.get("labelIds", [])
    is_draft = "DRAFT" in label_ids
    draft_note = " (This is a draft message)" if is_draft else ""

    headers = payload.get("headers", [])
    subject = next((h['value'] for h in headers if h['name'] == 'Subject'), "")
    to = next((h['value'] for h in headers if h['name'] == 'To'), "")
    cc = next((h['value'] for h in headers if h['name'] == 'Cc'), "")
    bcc = next((h['value'] for h in headers if h['name'] == 'Bcc'), "")
    sender = next((h['value'] for h in headers if h['name'] == 'From'), "Unknown")
    all_recipients = extract_emails(to) + extract_emails(cc) + extract_emails(bcc)
    
    body = get_email_body(payload)
    
    threat_score = 0
    severity = "low"
    flags = []
    attachment_names = []
    
    # Normalize the current user's full email
    user_email = username.lower() + ".zerotrace@gmail.com"

    # Check if the current user is a receiver
    is_receiver = user_email in all_recipients


    #Extract sender email(s)
    sender_emails = extract_emails(sender)

    # Ensure we only proceed if sender is found
    if sender_emails:
        sender_email = sender_emails[0].lower()
        
        sender_username = sender_email.split('@')[0]  # Extract the username part before '@'
        
        # Check if this is the current user's email
        is_sender = username.lower() + ".zerotrace@gmail.com" == sender_email
        
    else:
        sender_email = "unknown"
        
    
    confidential_found = False
    filename = ""
    file_hash = ""
    file_text = ""

    has_attachment = 'parts' in payload
    parts = payload.get('parts', []) if has_attachment else []

    if has_attachment:
        for part in parts:
            if 'filename' in part and part['filename']:
                filename = part['filename']
                attachment_names.append(filename)

                # 🔐 First, handle zip files
                zip_handled, zip_severity, zip_intent, zip_filename, zip_hash = process_zip_attachment(
                    service, message_id, part, sender, all_recipients, zip_file_history
                )

                if zip_handled:
                    flags.append("zip_attachment")
                    threat_score += 2  # You can customize this
                    update_history(username, "zip_attachment")
                    log_alert(
                        username,
                        f"Zip file attachment detected: {zip_filename}, hash: {zip_hash}",
                        zip_severity,
                        zip_intent
                    )
                    continue  # Skip further checks for this part (zip already handled)

            # 📄 Then handle other confidential file types
            confidential_found, filename, file_hash, file_text = attachment_contains_confidential_keywords(
                service, message_id, [part]
            )

            if confidential_found:
                threat_score += 2
                if is_sender : 
                    flags.append("sender_confidential_attachment")
                else : 
                    flags.append("recv_confidential_attachment")
                print(f"\033[91m[HIGH][INTENTIONAL] {sender}: SUSPICIOUS EMAIL!! (Subject: {subject}, To: {all_recipients} Attachments: {filename}{draft_note}")
                print(f"[ATTACHMENT CONTENT: {filename}]\n{file_text}\033[0m")


    # Only flag normal attachments if sent externally
    if has_attachment and not confidential_found:
        is_internal_to_internal = ORG_DOMAIN in sender and ORG_DOMAIN in all_recipients
        if is_internal_to_internal:
            threat_score += 0  # or 1 if you want to reduce severity
            if is_sender : 
                flags.append("sender_normal_attachment")
            else : 
                flags.append("recv_normal_attachment")
            

    if is_confidential(subject) or is_confidential(body) :
        threat_score += 2
        if is_sender : 
            flags.append("sender_confidential")
        else : 
            flags.append("recv_confidential")
    if is_any_external(all_recipients):
        threat_score += 2
        flags.append("sender_external")
    if contains_phishing_keywords(subject) or contains_phishing_keywords(body) :
        threat_score += 2
        if is_sender : 
            flags.append("sender_phishing")
        elif is_receiver : 
            flags.append("recv_phishing")
        else : 
            flags.append("phishing")
    if is_admin(username):
        threat_score += 1
        flags.append("admin_action")
    if "normal_attachment" in flags:
        threat_score += 1  
    if detect_spoofing(sender, headers):
        threat_score += 1
        flags.append("spoofing_detected")
    if detect_obfuscated_phishing(body):
        flags.append("phishing_obfuscated")


    if threat_score >= 6:
        severity = "critical"
    elif threat_score >= 5:
        severity = "high"
    elif threat_score >= 2:
        severity = "medium"

    if flags:
        update_history(username, ",".join(flags))
        intent = classify_intention(flags, is_any_external(to), username, subject, sender, all_recipients)
        attachment_info = f" Attachments: {', '.join(attachment_names)}" if attachment_names else ""
        log_alert(username, f"EMAIL DETECTED! (Subject: {subject}, To: {all_recipients}{draft_note}", severity, intent)

def check_user_inbox(username):
    try:
        service = get_gmail_service(username)
        results = service.users().messages().list(userId='me', maxResults=5).execute()
        messages = results.get('messages', [])
        for msg in messages:
            msg_data = service.users().messages().get(userId='me', id=msg['id'], format='full').execute()
            msg_data['id'] = msg['id']
            analyze_email_content(username, msg_data, service, msg['id'])

    except Exception as e:
        log_alert(username, f"Error checking inbox: {e}", "high")

def main():
    print("[*] Insider Threat Tool is now running. Press Ctrl+C to stop.")
    try:
        while True:
            for username in USERS:
                check_user_inbox(username)
            time.sleep(10)
    except KeyboardInterrupt:
        print("\n[!] Insider Threat Tool stopped by user.")

if __name__ == '__main__':
    main()







